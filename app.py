import streamlit as st
import logging
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
import os
import yaml
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import markdown
import re

# Configuração inicial da página
st.set_page_config(
    page_title="PsicoIA Pro - Relatórios Psicológicos",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Dicionário de usuários
USERS_DB = {
    'gabriel@aperdata.com': {'password': 'gabriel123', 'name': 'Administrador'},
    'maria@psicoiapro.com': {'password': 'maria123', 'name': 'Maria Silva'},
    'joao@psicoiapro.com': {'password': 'joao123', 'name': 'João Santos'}
}

# Constantes e configurações
REPORT_TYPES = {
    "Relatório de Devolutiva": "devolutiva",
    "Relatório de Evolução": "evolucao",
    "Relatório de Anamnese": "anamnese",
    "Relatório de Avaliação Psicológica Inicial": "avaliacao_inicial",
    "Relatório de Alta Terapêutica": "alta",
    "Relatório de Avaliação de Personalidade": "personalidade",
    "Relatório de Avaliação Neuropsicológica": "neuropsicologica",
    "Relatório de Acompanhamento Terapêutico": "acompanhamento",
    "Relatório de Intervenção Comportamental": "intervencao",
    "Relatório de Diagnóstico Psicológico": "diagnostico",
    "Relatório de Avaliação Emocional": "emocional",
    "Relatório para Escolas": "escolar",
    "Relatório de Avaliação Infantil": "infantil",
    "Relatório de Avaliação para Orientação Profissional": "profissional",
    "Relatório de Avaliação Familiar": "familiar",
    "Relatório de Sessão Terapêutica": "sessao",
    "Relatório de Feedback para o Paciente e Família": "feedback"
}

TONE_DESCRIPTIONS = {
    "Tom Formal e Técnico": "Use linguagem técnica e profissional, priorizando termos científicos e mantendo um tom objetivo e formal.",
    "Tom Acessível e Didático": "Use linguagem clara e acessível, explicando conceitos técnicos de forma didática e compreensível.",
    "Tom Colaborativo e Empático": "Use linguagem acolhedora e empática, mantendo o profissionalismo mas priorizando a conexão humana."
}

ABORDAGENS_TERAPEUTICAS = [
    'Terapia Cognitivo-Comportamental',
    'Psicanálise',
    'Terapia Humanista',
    'Terapia Sistêmica',
    'Terapia Integrativa',
    'Terapia ABA',
    'Gestalt-terapia',
    'Terapia Analítica',
    'Terapia Centrada na Pessoa',
    'Terapia Comportamental'
]

GENEROS = ['Masculino', 'Feminino', 'Não-binário', 'Prefiro não especificar']

# CSS Customizado
st.markdown("""
<style>
.main-header {
    font-size: 2.5rem;
    color: #2c3e50;
    text-align: center;
    margin-bottom: 2rem;
}
.section-header {
    color: #34495e;
    padding: 1rem 0;
    border-bottom: 2px solid #ecf0f1;
}
.stButton>button {
    background-color: #3498db;
    color: white;
    font-weight: bold;
    padding: 0.75rem 1.5rem;
    border-radius: 5px;
    border: none;
    transition: all 0.3s ease;
}
.stButton>button:hover {
    background-color: #2980b9;
    box-shadow: 0 2px 4px rgba(0, 0, 0, 0.2);
}
.custom-info-box {
    background-color: #f8f9fa;
    padding: 1rem;
    border-radius: 5px;
    border-left: 5px solid #3498db;
    margin: 1rem 0;
}
</style>
""", unsafe_allow_html=True)

def check_login(username, password):
    """Verifica as credenciais de login"""
    if username in USERS_DB:
        if USERS_DB[username]['password'] == password:
            return True, USERS_DB[username]['name']
    return False, None

def convert_markdown_to_docx(markdown_text):
    """Converte texto markdown para arquivo DOCX"""
    doc = Document()
    
    header = doc.add_heading('Relatório Psicológico', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cleaned_text = re.sub(r'\*{1,2}', '', markdown_text)
    cleaned_text = re.sub(r'^#+\s', '', cleaned_text, flags=re.MULTILINE)
    cleaned_text = re.sub(r'^\-\s', '• ', cleaned_text, flags=re.MULTILINE)
    
    paragraphs = cleaned_text.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            p.add_run(para.strip())
            p.paragraph_format.space_after = Pt(12)
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def get_specific_fields(report_type):
    """Retorna campos específicos baseados no tipo de relatório"""
    fields = {}
    
    if report_type == "devolutiva":
        fields.update({
            "resultados_avaliacao": st.text_area("Resultados da Avaliação:", height=150),
            "interpretacao": st.text_area("Interpretação dos Resultados:", height=150),
            "recomendacoes": st.text_area("Recomendações:", height=150),
            "recursos_utilizados": st.text_area("Recursos e Testes Utilizados:", height=100)
        })
    
    elif report_type == "evolucao":
        fields.update({
            "periodo_avaliado": st.text_input("Período Avaliado:"),
            "objetivos_terapeuticos": st.text_area("Objetivos Terapêuticos:", height=150),
            "progresso": st.text_area("Progresso Observado:", height=150),
            "desafios": st.text_area("Desafios Encontrados:", height=150),
            "estrategias": st.text_area("Estratégias Utilizadas:", height=150)
        })
    
    elif report_type == "anamnese":
        fields.update({
            "queixa_principal": st.text_area("Queixa Principal:", height=150),
            "historico_sintomas": st.text_area("Histórico dos Sintomas:", height=150),
            "historico_familiar": st.text_area("Histórico Familiar:", height=150),
            "historico_medico": st.text_area("Histórico Médico:", height=150),
            "desenvolvimento": st.text_area("História do Desenvolvimento:", height=150)
        })
    
    elif report_type == "avaliacao_inicial":
        fields.update({
            "demanda": st.text_area("Demanda Inicial:", height=150),
            "sintomas_atuais": st.text_area("Sintomas Atuais:", height=150),
            "historico_tratamentos": st.text_area("Histórico de Tratamentos:", height=150),
            "suporte_social": st.text_area("Rede de Suporte Social:", height=150)
        })
    
    elif report_type == "alta":
        fields.update({
            "motivo_alta": st.text_area("Motivo da Alta:", height=150),
            "objetivos_alcancados": st.text_area("Objetivos Alcançados:", height=150),
            "progresso_final": st.text_area("Progresso Final:", height=150),
            "recomendacoes_futuras": st.text_area("Recomendações Futuras:", height=150)
        })
    
    elif report_type == "personalidade":
        fields.update({
            "instrumentos_utilizados": st.text_area("Instrumentos de Avaliação Utilizados:", height=150),
            "resultados_personalidade": st.text_area("Resultados da Avaliação de Personalidade:", height=150),
            "perfil_psicologico": st.text_area("Perfil Psicológico:", height=150),
            "implicacoes_praticas": st.text_area("Implicações Práticas:", height=150)
        })
    
    elif report_type == "neuropsicologica":
        fields.update({
            "funcoes_avaliadas": st.text_area("Funções Cognitivas Avaliadas:", height=150),
            "instrumentos_neuropsicologicos": st.text_area("Instrumentos Neuropsicológicos Utilizados:", height=150),
            "resultados_cognitivos": st.text_area("Resultados por Função Cognitiva:", height=150),
            "conclusao_diagnostica": st.text_area("Conclusão Diagnóstica:", height=150),
            "recomendacoes_reabilitacao": st.text_area("Recomendações para Reabilitação:", height=150)
        })
    
    elif report_type == "acompanhamento":
        fields.update({
            "periodo_acompanhamento": st.text_input("Período de Acompanhamento:"),
            "objetivos_alcancados": st.text_area("Objetivos Alcançados:", height=150),
            "evolucao_observada": st.text_area("Evolução Observada:", height=150),
            "aspectos_relevantes": st.text_area("Aspectos Relevantes:", height=150),
            "proximos_passos": st.text_area("Próximos Passos:", height=150)
        })
    
    elif report_type == "intervencao":
        fields.update({
            "comportamentos_alvo": st.text_area("Comportamentos-Alvo:", height=150),
            "estrategias_intervencao": st.text_area("Estratégias de Intervenção:", height=150),
            "resultados_obtidos": st.text_area("Resultados Obtidos:", height=150),
            "ajustes_necessarios": st.text_area("Ajustes Necessários:", height=150)
        })
    
    elif report_type == "diagnostico":
        fields.update({
            "sintomas_apresentados": st.text_area("Sintomas Apresentados:", height=150),
            "criterios_diagnosticos": st.text_area("Critérios Diagnósticos:", height=150),
            "diagnostico_diferencial": st.text_area("Diagnóstico Diferencial:", height=150),
            "conclusao_diagnostica": st.text_area("Conclusão Diagnóstica:", height=150),
            "plano_tratamento": st.text_area("Plano de Tratamento:", height=150)
        })
    
    elif report_type == "emocional":
        fields.update({
            "estado_emocional": st.text_area("Estado Emocional Atual:", height=150),
            "fatores_estresse": st.text_area("Fatores de Estresse:", height=150),
            "recursos_enfrentamento": st.text_area("Recursos de Enfrentamento:", height=150),
            "suporte_social": st.text_area("Suporte Social:", height=150),
            "recomendacoes": st.text_area("Recomendações:", height=150)
        })
    
    elif report_type == "escolar":
        fields.update({
            "desempenho_academico": st.text_area("Desempenho Acadêmico:", height=150),
            "comportamento_escolar": st.text_area("Comportamento em Ambiente Escolar:", height=150),
            "relacoes_interpessoais": st.text_area("Relações Interpessoais:", height=150),
            "necessidades_especificas": st.text_area("Necessidades Específicas:", height=150),
            "recomendacoes_escola": st.text_area("Recomendações para a Escola:", height=150)
        })
    
    elif report_type == "infantil":
        fields.update({
            "desenvolvimento_atual": st.text_area("Desenvolvimento Atual:", height=150),
            "comportamento_observado": st.text_area("Comportamento Observado:", height=150),
            "interacao_social": st.text_area("Interação Social:", height=150),
            "aspectos_familiares": st.text_area("Aspectos Familiares:", height=150),
            "recomendacoes_pais": st.text_area("Recomendações aos Pais:", height=150)
        })
    
    elif report_type == "profissional":
        fields.update({
            "interesses_profissionais": st.text_area("Interesses Profissionais:", height=150),
            "habilidades_identificadas": st.text_area("Habilidades Identificadas:", height=150),
            "valores_trabalho": st.text_area("Valores Relacionados ao Trabalho:", height=150),
            "areas_recomendadas": st.text_area("Áreas Recomendadas:", height=150),
            "plano_desenvolvimento": st.text_area("Plano de Desenvolvimento:", height=150)
        })
    
    elif report_type == "familiar":
        fields.update({
            "dinamica_familiar": st.text_area("Dinâmica Familiar:", height=150),
            "padroes_relacionamento": st.text_area("Padrões de Relacionamento:", height=150),
            "conflitos_identificados": st.text_area("Conflitos Identificados:", height=150),
            "recursos_familiares": st.text_area("Recursos Familiares:", height=150),
            "recomendacoes_familia": st.text_area("Recomendações para a Família:", height=150)
        })
    
    elif report_type == "sessao":
        fields.update({
            "temas_abordados": st.text_area("Temas Abordados:", height=150),
            "tecnicas_utilizadas": st.text_area("Técnicas Utilizadas:", height=150),
            "respostas_paciente": st.text_area("Respostas do Paciente:", height=150),
            "insights_obtidos": st.text_area("Insights Obtidos:", height=150),
            "planejamento_proxima": st.text_area("Planejamento para Próxima Sessão:", height=150)
        })
    
    elif report_type == "feedback":
        fields.update({
            "progresso_observado": st.text_area("Progresso Observado:", height=150),
            "pontos_positivos": st.text_area("Pontos Positivos:", height=150),
            "areas_desenvolvimento": st.text_area("Áreas para Desenvolvimento:", height=150),
            "orientacoes_praticas": st.text_area("Orientações Práticas:", height=150),
            "proximos_objetivos": st.text_area("Próximos Objetivos:", height=150)
        })
    
    return fields

def create_prompt(report_type, tone, patient_data, specific_fields):
    """Cria o prompt para o modelo de IA"""
    base_template = f"""
    Você é um assistente especializado em psicologia, focado na geração de {report_type}.
    
    Tom do relatório: {tone}
    
    Dados do paciente:
    Nome: {patient_data['nome']}
    Idade: {patient_data['idade']} anos
    Gênero: {patient_data['genero']}
    Data da avaliação: {patient_data['data_avaliacao']}
    Abordagem terapêutica: {patient_data['abordagem_terapeutica']}

    Informações específicas:
    """
    
    for key, value in specific_fields.items():
        base_template += f"{key}: {value}\n"
    
    base_template += "\nPor favor, gere um relatório profissional e detalhado."
    
    return base_template

def main():
    # Inicializa o estado da sessão
    if 'logged_in' not in st.session_state:
        st.session_state['logged_in'] = False

    # Tratamento do login
    if not st.session_state['logged_in']:
        st.title("Login - PsicoIA Pro")
        
        username = st.text_input("Email")
        password = st.text_input("Senha", type="password")
        
        if st.button("Entrar"):
            if username and password:
                success, user_name = check_login(username, password)
                if success:
                    st.session_state['logged_in'] = True
                    st.session_state['username'] = username
                    st.session_state['user_name'] = user_name
                    st.success(f"Bem-vindo(a), {user_name}!")
                    st.rerun()
                else:
                    st.error("Email ou senha incorretos!")
            else:
                st.warning("Por favor, preencha todos os campos!")
        return

    # Aplicação principal (mostrada apenas quando logado)
    with st.sidebar:
        st.title("PsicoIA Pro")
        st.subheader("AperData Solutions")
        st.markdown("""
        Entre em contato:
    
        🌐 [aperdata.com](https://aperdata.com)  
        📱 WhatsApp: 11 98854-3437  
        📧 Email: [gabriel@aperdata.com](mailto:gabriel@aperdata.com)
        """)
        if st.button("Logout"):
            st.session_state['logged_in'] = False
            st.rerun()
        
        menu_selection = st.radio(
            "Menu Principal",
            ["Gerar Relatório", "Configurações", "Sobre"]
        )

    if menu_selection == "Gerar Relatório":
        st.title("Gerador de Relatórios Psicológicos")

        col1, col2 = st.columns(2)
        with col1:
            report_type = st.selectbox(
                "Selecione o tipo de relatório:",
                list(REPORT_TYPES.keys())
            )

        with col2:
            tone = st.selectbox(
                "Selecione o tom do relatório:",
                list(TONE_DESCRIPTIONS.keys())
            )

        tab1, tab2, tab3 = st.tabs([
            "Dados do Paciente",
            "Informações Específicas",
            "Gerar Relatório"
        ])

        with tab1:
            col1, col2 = st.columns(2)
            with col1:
                nome = st.text_input("Nome completo do paciente:")
                idade = st.number_input("Idade:", min_value=0, max_value=120)
                genero = st.selectbox("Gênero:", GENEROS)

            with col2:
                data_avaliacao = st.date_input("Data da avaliação:")
                abordagem_terapeutica = st.selectbox(
                    "Abordagem terapêutica:",
                    ABORDAGENS_TERAPEUTICAS
                )

        with tab2:
            specific_fields = get_specific_fields(REPORT_TYPES[report_type])

        with tab3:
            if st.button("Gerar Relatório"):
                try:
                    # Usa st.secrets em vez de config.yaml
                    api_key = st.secrets["OPENAI_API_KEY"]
                    os.environ['OPENAI_API_KEY'] = api_key

                    model = ChatOpenAI(
                        model_name='gpt-4-turbo',
                        temperature=0.7
                    )

                    patient_data = {
                        "nome": nome,
                        "idade": idade,
                        "genero": genero,
                        "data_avaliacao": data_avaliacao.strftime("%d/%m/%Y"),
                        "abordagem_terapeutica": abordagem_terapeutica
                    }
                    
                    prompt = create_prompt(report_type, tone, patient_data, specific_fields)
                    
                    with st.spinner("Gerando relatório..."):
                        response = model.invoke(prompt)
                        docx_file = convert_markdown_to_docx(response.content)
                        
                        st.success("Relatório gerado com sucesso!")
                        st.markdown(response.content)

                        st.download_button(
                            "Download Relatório (DOCX)",
                            docx_file,
                            f"relatorio_{REPORT_TYPES[report_type]}_{datetime.now().strftime('%Y%m%d')}.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Erro ao gerar relatório: {str(e)}")

    elif menu_selection == "Configurações":
        st.title("Configurações")
        st.write("Área em desenvolvimento...")

    elif menu_selection == "Sobre":
        st.title("Sobre o PsicoIA Pro")
        st.markdown("""
        O PsicoIA Pro é uma ferramenta avançada de inteligência artificial desenvolvida 
        para otimizar e elevar a qualidade da geração de relatórios psicológicos.
        
        ### Recursos
        - 17 tipos diferentes de relatórios psicológicos
        - Seleção de tom para personalização
        - Interface intuitiva e profissional
        - Integração com IA avançada
        - Formato padronizado seguindo normas técnicas
        """)

if __name__ == "__main__":
    main()                                               
